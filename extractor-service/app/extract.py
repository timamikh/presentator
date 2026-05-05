"""Pure extraction helpers — no I/O over HTTP, no DB.

Each function takes raw bytes (or a path) and returns plain text. Keeping
extraction pure-functional makes unit-testing trivial (see tests/test_extract.py)
and lets the service handle PDF / DOCX / TXT through a single dispatcher.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

WHITESPACE_RE = re.compile(r"[ \t]+")
NEWLINES_RE = re.compile(r"\n{3,}")


@dataclass(frozen=True)
class ExtractionResult:
    """Output of any extractor.

    text: raw text from the document, normalized whitespace.
    summary: short summary (1-3 paragraphs). For now derived from the text head;
             a future iteration will plug in an LLM call here.
    width / height: only set for image attachments.
    """

    text: str
    summary: str
    width: Optional[int] = None
    height: Optional[int] = None


def _normalize_text(text: str, limit: int) -> str:
    if not text:
        return ""
    text = WHITESPACE_RE.sub(" ", text)
    text = NEWLINES_RE.sub("\n\n", text)
    text = text.strip()
    if len(text) > limit:
        text = text[:limit]
    return text


def _summarize_head(text: str, max_chars: int) -> str:
    """Naive heuristic summary: first paragraph(s) until max_chars.

    LLM-based summarization is a future hook (skill: replace this body and add
    httpx call to LLM with system prompt 'condense to 3 paragraphs'). The
    pipeline already passes content_summary downstream, so the upgrade is
    transparent to consumers.
    """
    if not text:
        return ""
    snippet = text.strip()[:max_chars]
    last_period = snippet.rfind(". ")
    if last_period > max_chars // 2:
        snippet = snippet[: last_period + 1]
    return snippet


def extract_pdf(data: bytes, *, text_limit: int, summary_max_chars: int) -> ExtractionResult:
    from pdfminer.high_level import extract_text

    raw = extract_text(io.BytesIO(data)) or ""
    text = _normalize_text(raw, text_limit)
    return ExtractionResult(text=text, summary=_summarize_head(text, summary_max_chars))


def extract_docx(data: bytes, *, text_limit: int, summary_max_chars: int) -> ExtractionResult:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                parts.append(" | ".join(cells))
    raw = "\n".join(parts)
    text = _normalize_text(raw, text_limit)
    return ExtractionResult(text=text, summary=_summarize_head(text, summary_max_chars))


def extract_txt(data: bytes, *, text_limit: int, summary_max_chars: int) -> ExtractionResult:
    try:
        raw = data.decode("utf-8")
    except UnicodeDecodeError:
        raw = data.decode("utf-8", errors="replace")
    text = _normalize_text(raw, text_limit)
    return ExtractionResult(text=text, summary=_summarize_head(text, summary_max_chars))


def extract_image_metadata(data: bytes) -> ExtractionResult:
    """Pillow-only metadata extraction: width/height, no OCR (kept out of scope)."""
    from PIL import Image

    with Image.open(io.BytesIO(data)) as img:
        return ExtractionResult(
            text="",
            summary="",
            width=img.width,
            height=img.height,
        )


def dispatch(
    *,
    path: Path,
    mime_type: Optional[str],
    text_limit: int,
    summary_max_chars: int,
) -> ExtractionResult:
    """Pick an extractor by mime type with file-extension fallback."""
    data = path.read_bytes()
    mt = (mime_type or "").lower()
    suffix = path.suffix.lower()

    if mt.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}:
        return extract_image_metadata(data)

    if "pdf" in mt or suffix == ".pdf":
        return extract_pdf(data, text_limit=text_limit, summary_max_chars=summary_max_chars)

    if "officedocument.wordprocessingml" in mt or suffix == ".docx":
        return extract_docx(data, text_limit=text_limit, summary_max_chars=summary_max_chars)

    if mt.startswith("text/") or suffix in {".txt", ".md", ".csv", ".json", ".xml"}:
        return extract_txt(data, text_limit=text_limit, summary_max_chars=summary_max_chars)

    logger.warning("dispatch: unsupported mime=%s suffix=%s for %s", mt, suffix, path)
    return ExtractionResult(text="", summary="")
