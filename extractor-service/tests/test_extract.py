"""Unit tests for pure extraction helpers (no I/O over HTTP)."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from app.extract import (
    ExtractionResult,
    dispatch,
    extract_docx,
    extract_image_metadata,
    extract_pdf,
    extract_txt,
)


def test_extract_txt_normalizes_whitespace_and_truncates():
    raw = b"hello   world\n\n\n\nnext line"
    result = extract_txt(raw, text_limit=100, summary_max_chars=20)
    assert "hello world" in result.text
    assert result.text.count("\n\n\n") == 0
    assert len(result.summary) <= 20 + 1


def test_extract_txt_handles_invalid_utf8():
    raw = b"\xff\xfe\x00bad" + "хорошо".encode("utf-8")
    result = extract_txt(raw, text_limit=100, summary_max_chars=100)
    assert "хорошо" in result.text


def test_extract_txt_truncates_at_text_limit():
    raw = ("a" * 1000).encode("utf-8")
    result = extract_txt(raw, text_limit=100, summary_max_chars=50)
    assert len(result.text) == 100


def test_extract_pdf_minimal_document(tmp_path: Path):
    """Real PDF parsing — uses pdfminer.six on a tiny in-memory PDF."""
    try:
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed (used only for synthetic PDF test)")

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Hello PDF World")
    c.save()
    pdf_bytes = buf.getvalue()

    result = extract_pdf(pdf_bytes, text_limit=10000, summary_max_chars=200)
    assert "Hello PDF World" in result.text


def test_extract_docx_with_paragraphs_and_tables():
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "value"

    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    result = extract_docx(docx_bytes, text_limit=10000, summary_max_chars=200)
    assert "First paragraph." in result.text
    assert "Second paragraph." in result.text
    assert "name | value" in result.text


def test_extract_image_metadata_returns_dimensions():
    try:
        from PIL import Image
    except ImportError:
        pytest.skip("Pillow not installed")

    buf = io.BytesIO()
    Image.new("RGB", (320, 200), color="red").save(buf, format="PNG")
    img_bytes = buf.getvalue()

    result = extract_image_metadata(img_bytes)
    assert result.width == 320
    assert result.height == 200
    assert result.text == ""


def test_dispatch_picks_txt_for_text_mime(tmp_path: Path):
    f = tmp_path / "note.txt"
    f.write_bytes(b"plain text content")
    result = dispatch(
        path=f,
        mime_type="text/plain",
        text_limit=1000,
        summary_max_chars=100,
    )
    assert "plain text content" in result.text


def test_dispatch_falls_back_to_extension(tmp_path: Path):
    f = tmp_path / "note.txt"
    f.write_bytes(b"some content")
    result = dispatch(
        path=f,
        mime_type="application/octet-stream",
        text_limit=1000,
        summary_max_chars=100,
    )
    assert "some content" in result.text


def test_dispatch_returns_empty_for_unsupported(tmp_path: Path):
    f = tmp_path / "binary.dat"
    f.write_bytes(b"\x00\x01\x02\x03")
    result = dispatch(
        path=f,
        mime_type="application/octet-stream",
        text_limit=1000,
        summary_max_chars=100,
    )
    assert result.text == ""
    assert result.summary == ""
    assert result.width is None
    assert result.height is None


def test_summary_does_not_exceed_max_chars():
    raw = ("Sentence one. " * 200).encode("utf-8")
    result = extract_txt(raw, text_limit=100000, summary_max_chars=50)
    assert len(result.summary) <= 50 + 1
